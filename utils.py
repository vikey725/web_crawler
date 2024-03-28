'''The utils Module'''

from __future__ import annotations

import os
import re

import docx
from docx.shared import Inches
import requests

from configs import Configs

class Utils:
    """Module to store utility functions for WebCrawler"""

    def __init__(self):
        pass

    @staticmethod
    def add_hyperlink(paragraph, text, url):
        """_summary_

        Args:
            paragraph (_type_): _description_
            text (_type_): _description_
            url (_type_): _description_

        Returns:
            _type_: _description_
        """
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a new run object (a wrapper over a 'w:r' element)
        new_run = docx.text.run.Run(
            docx.oxml.shared.OxmlElement('w:r'), paragraph)
        new_run.text = text

        new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        new_run.font.underline = True

        # Join all the xml elements together
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)

        return hyperlink


    def download_image(self, image_url):
        image_name = image_url.split('/')[-1]
        if not os.path.exists(Configs.IMAGE_DIR):
            os.mkdir(Configs.IMAGE_DIR)
        image_path = os.path.join(Configs.IMAGE_DIR, image_name)
        try:
            data = requests.get(image_url).content
            with open(image_path, 'wb') as f:
                f.write(data) 
            return image_path
        except Exception as e:
            print(e, image_url)
            return None


    def get_file_path(self, page_url):
        page_url = page_url[8:]
        if not os.path.exists(Configs.DATA_DIR):
            os.mkdir(Configs.DATA_DIR)
        return os.path.join(Configs.DATA_DIR, '-'.join(page_url.split('/')) + '.docx')

            
    def write_docs(self, content, page_url):
        file_path = self.get_file_path(page_url)
        document = docx.Document()
        for data in content:
            if data['tag_name'] == 'table':
                table_data = data['table_data']
                num_rows, num_cols = len(table_data), len(table_data[0])
                table = document.add_table(rows=num_rows, cols=num_cols)
                for i in range(num_rows):
                    cells = table.rows[i].cells
                    for j in range(num_cols):
                        cells[j].text = table_data[i][j]
            elif data['tag_name'] == 'a':
                paragraph = document.add_paragraph()
                Utils.add_hyperlink(paragraph, data.get('text'), data.get('link'))
            elif data['tag_name'] == 'img':
                image_url = self.preprocess_image_url(data.get('img_src'))
                image_path = self.download_image(image_url)
                if image_path:
                    if image_path.find('Icon.png') != -1:
                        document.add_picture(image_path, width=Inches(0.25))
                    else:
                        document.add_picture(image_path, width=Inches(4))
                else:
                    paragraph = document.add_paragraph('Image not Found')
                    paragraph.add_run('italic.').italic = True
            elif data['tag_name'] in set(['h1', 'h2', 'h3']):
                document.add_heading(data.get('text'), 0)
            else:
                document.add_paragraph(data.get('text'))
        document.save(file_path)


    def preprocess_image_url(self, url):
        regex = r'(\S*\.png|\S*\.jpg)\S*'
        try:
            matches = re.findall(regex, url)
            return matches[0]
        except IndexError as e:
            print(e, url)
            return url


    def preprocess_web_url(self, url):
        regex = r'(\S*\d+)\S*'
        try:
            matches = re.findall(regex, url)
            return matches[0]
        except IndexError as e:
            print(e, 'url:', url)
            return url


if __name__ == '__main__':
    utils = Utils()
    utils.preprocess_image_url('https://docs/aryaka.com/download/attachments/1542430/UntticketingIcon.png?version=1&modificationDate=1711559163535&cacheVersion=1&api=v2')
